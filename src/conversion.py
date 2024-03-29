# 以下格式标准均来自该文档：
#（整本模板-普招）信息与智能工程学院2020级毕业论文（设计）（更新：2023.12）

import json
import math
import os
import os.path as path
import pprint
import re
import shutil
import sys
import time

import linkify_it
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
# WD_LINE_SPACING 为行距。
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_PARAGRAPH_ALIGNMENT, WD_TAB_LEADER,
                            WD_UNDERLINE)
from docx.oxml import CT_Inline, OxmlElement, ns, xmlchemy
# 颜色、单位。
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import CharacterStyle, ParagraphStyle
from docxcompose.composer import Composer
from markdown_it import MarkdownIt
from mdit_py_plugins import front_matter
from requests import get as requests_get
from yaml import Loader as yaml_Loader
from yaml import load as yaml_load

from myconfig import load_config
from myenums import CNFontFaces, FontSizes, bcolors
from mynode_server import start_node_server
from myutils import (add_page_number, add_toc, calculate_char_map_space_sum,
                     count_char_map, create_attribute, create_element,
                     create_page_number_type, get_count_char_space_width,
                     list_number, preprint_content_in_map, print_faild,
                     set_bold, set_font_face, set_font_face_en)
from obsdian_image_plugin import obsdian_image_plugin

cfg = None
try:
  cfg = load_config(sys.argv[1] if sys.argv.__len__() > 1 else None)
except Exception as e:
  print_faild("解析配置文件时出错。", e)
  exit(0)

src_path = cfg["src_path"]
target_path = cfg["target_path"]
obsdian_image_find_paths = cfg["obsdian_image_find_paths"]
template_path = cfg["template_path"]
temp_dir_path = cfg["temp_dir_path"]
nodejs_path = cfg["nodejs_path"]


temp_dir_counter = 0

def get_id():
  global temp_dir_counter
  temp_dir_counter = temp_dir_counter + 1
  return temp_dir_counter


if path.exists(temp_dir_path):
  shutil.rmtree(temp_dir_path)
os.mkdir(temp_dir_path)

node_server = start_node_server(nodejs_path)

md = MarkdownIt("gfm-like", {
  "html": True,
}).enable('table').use(
  front_matter.index.front_matter_plugin
  ).use(
    obsdian_image_plugin,
    options = { 
      "obsdian_image_find_paths": obsdian_image_find_paths
    }
  )


# 页眉的格式设置函数
def procress_header(header):
  header.paragraphs[0].style.font.size = Pt(FontSizes.小五) # type: ignore
  header.paragraphs[0].style.paragraph_format.first_line_indent = False # type: ignore
  set_font_face(header.paragraphs[0].style, CNFontFaces.宋体, 'Times New Roman') # type: ignore
  header.is_linked_to_previous = False
  
  # 为 Header 添加下边框
  element: xmlchemy.BaseOxmlElement = header._element
  wp: xmlchemy.BaseOxmlElement = element.first_child_found_in("w:p") # type: ignore
  wpr: xmlchemy.BaseOxmlElement = wp.first_child_found_in("w:pPr") # type: ignore
  pbdr = create_element("w:pBdr")
  pbdr_bottom = create_element("w:bottom")
  create_attribute(pbdr_bottom, "w:val", "single")
  create_attribute(pbdr_bottom, "w:sz", "4")
  create_attribute(pbdr_bottom, "w:space", "1")
  create_attribute(pbdr_bottom, "w:color", "auto")
  pbdr.append(pbdr_bottom)
  wpr.append(pbdr)

def get_meta_data_key(key_name: str):
  global meta_data
  result = ""
  try:
    result = meta_data[key_name]
  except:
    result = "（请手动填写）"
    print(
      f"{bcolors.WARNING}[信息缺失] 你没有提供{key_name}，请在生成后手动填写，或在 Markdown 文档开头使用设置 {key_name} 属性。\n{bcolors.ENDC}"
    )
  return result

doc: DocumentType = Document()
meta_data = {}
composer = Composer(doc)
src_lines = []
tokens = []

# 设置“正文”样式的具体样式。
# 
normal_style: ParagraphStyle = doc.styles["Normal"]  # type: ignore
normal_style.font.size = Pt(FontSizes.小四)
# 中文字体：小四号宋体；英文字体：一律为 Times New Roman。
set_font_face(normal_style, CNFontFaces.宋体, 'Times New Roman')
# 1.5倍行距
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
# 正文段落要统一首行缩进2个字符
normal_style.paragraph_format.first_line_indent = Pt(FontSizes.小四 * 2)


# 设置“各级标题”样式的具体样式。
# 一级标题：黑体，加粗，小三，居中，1.5倍行距，段前1行，段后2行。
# 二级标题：黑体，四号，1.5倍行距，段前1行，段后1行
# 三级标题：黑体，四号，1.5倍行距，段前0.5行，段后0行
# 注意，四级以后的标题没有明确规定。因此这里按照合理的合理的方法编排字号。
heading_style_sizes = [
  FontSizes.小三,
  FontSizes.四号,
  FontSizes.四号,
  FontSizes.小四,
  FontSizes.五号,
  FontSizes.小五
]
heading_spacing_befores = [1, 1, 0.5, 0.5, 0.5, 0.5]
heading_spacing_afters = [2, 1, 0, 0, 0, 0]
for heading_level in range(1, 7):
  font_size = heading_style_sizes[heading_level - 1]
  spacing_before = heading_spacing_befores[heading_level - 1]
  spacing_after = heading_spacing_afters[heading_level - 1]


  # 由于 python docx 无法设置 `Heading x` 的 字体，因此在这里新建 Heading。
  # heading_style: ParagraphStyle = doc.styles.add_style(f'Heading {heading_level}', WD_STYLE_TYPE.PARAGRAPH) # type: ignore
  heading_style: ParagraphStyle = doc.styles[f"Heading {heading_level}"] # type: ignore
  heading_style.style_id = f"Heading{heading_level}"
  heading_style.hidden = False
  heading_style.quick_style = True

  heading_style.font.size = Pt(font_size)
  set_bold(heading_style)
  set_font_face(heading_style, CNFontFaces.黑体, 'Times New Roman')
  heading_style.font.color.rgb = RGBColor.from_string("000000")

  paragraph_format = heading_style.paragraph_format
  paragraph_format.first_line_indent = 0
  # 行距使用当前字号大小
  paragraph_format.space_before = Pt(font_size * spacing_before)
  paragraph_format.space_after = Pt(font_size * spacing_after)
  # 1.5倍行距
  paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
  
  # 每个章节前进行页面间隔
  if heading_level == 1:
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.page_break_before = True

# 代码样式
# 代码格式为五号，英文为Times New Roman；中文为宋体，单倍行距。
code_style: ParagraphStyle = doc.styles.add_style("代码", WD_STYLE_TYPE.PARAGRAPH, True) # type: ignore
code_style.font.size = Pt(FontSizes.五号)
code_style.quick_style = True
code_style.hidden = False
set_font_face(code_style, CNFontFaces.宋体, 'Times New Roman')
code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
code_style.paragraph_format.space_before = 0
code_style.paragraph_format.space_after = 0

# 图名
# 图名用五号，宋体，1.5倍行距，图名结束后不用标点。
img_desc_style: ParagraphStyle = doc.styles.add_style("插图标题", WD_STYLE_TYPE.PARAGRAPH, True) # type: ignore
img_desc_style.priority = 99
img_desc_style.quick_style = True
img_desc_style.hidden = False
img_desc_style.font.size = Pt(FontSizes.五号)
set_font_face(img_desc_style, CNFontFaces.宋体, 'Times New Roman')
img_desc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_desc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
img_desc_style.paragraph_format.space_before = 0
img_desc_style.paragraph_format.space_after = 0

# 参考文献样式
refs_style: ParagraphStyle = doc.styles.add_style("参考文献", WD_STYLE_TYPE.PARAGRAPH, True) # type: ignore
refs_style.priority = 99
refs_style.quick_style = True
refs_style.hidden = False
refs_style.font.size = Pt(FontSizes.小五)
set_font_face(refs_style, CNFontFaces.宋体, 'Times New Roman')
refs_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
refs_style.paragraph_format.space_before = Pt(FontSizes.小五)
refs_style.paragraph_format.space_after = Pt(FontSizes.小五)
refs_style.paragraph_format.first_line_indent = Cm(-0.76)
refs_style.paragraph_format.left_indent = Cm(0.76)

# 封面文本样式
cover_text_style: ParagraphStyle = doc.styles.add_style("封面文本", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
cover_text_style.priority = 99
cover_text_style.quick_style = True
cover_text_style.hidden = False
cover_text_style.font.size = Pt(FontSizes.小三)
set_font_face(cover_text_style, CNFontFaces.宋体, 'Times New Roman')
cover_text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_text_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
cover_text_style.paragraph_format.space_before = 0
cover_text_style.paragraph_format.space_after = 0
cover_text_style.paragraph_format.first_line_indent = Pt(0)

# 封面文本内容（字符）样式
cover_text_content_style: CharacterStyle =  doc.styles.add_style("封面文本内容", WD_STYLE_TYPE.CHARACTER, True) # type: ignore
cover_text_content_style.priority = 99
cover_text_content_style.quick_style = True
cover_text_content_style.hidden = False
cover_text_content_style.font.underline = WD_UNDERLINE.SINGLE

# 学术诚信声明正文
# 宋体，四号，行距：28
ais_text_style: ParagraphStyle = doc.styles.add_style("学术诚信声明正文", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
ais_text_style.priority = 99
ais_text_style.quick_style = True
ais_text_style.hidden = False
ais_text_style.font.size = Pt(FontSizes.四号)
set_font_face(ais_text_style, CNFontFaces.宋体, 'Times New Roman')
ais_text_style.paragraph_format.line_spacing = Pt(28)

# 摘要标题
# 三号宋体加粗，居中，1.5倍行距，段前1行、段后1行
abstract_title_style: ParagraphStyle = doc.styles.add_style("摘要标题", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
abstract_title_style.priority = 99
abstract_title_style.quick_style = True
abstract_title_style.hidden = False
abstract_title_style.font.size = Pt(FontSizes.三号)
set_font_face(abstract_title_style, CNFontFaces.宋体, 'Times New Roman')
set_bold(abstract_title_style)
abstract_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract_title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
abstract_title_style.paragraph_format.space_before = Pt(FontSizes.三号)
abstract_title_style.paragraph_format.space_after = Pt(FontSizes.三号)
abstract_title_style.paragraph_format.first_line_indent = 0

# 摘要内容
# 五号黑体并加方括号，左边顶格。五号楷体，1.5倍行距
abstract_content_style: ParagraphStyle = doc.styles.add_style("摘要内容", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
abstract_content_style.priority = 99
abstract_content_style.quick_style = True
abstract_content_style.hidden = False
abstract_content_style.font.size = Pt(FontSizes.五号)
abstract_content_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
abstract_content_style.paragraph_format.space_before = 0
abstract_content_style.paragraph_format.space_after = 0
abstract_content_style.paragraph_format.first_line_indent = 0

# 摘要内容标签
# 黑体
abstract_content_tag_style: CharacterStyle = doc.styles.add_style("摘要内容标签", WD_STYLE_TYPE.CHARACTER, False) # type: ignore
abstract_content_tag_style.priority = 99
abstract_content_style.quick_style = True
abstract_content_style.hidden = False
set_font_face(abstract_content_style, CNFontFaces.黑体, 'Times New Roman')

# 摘要内容正文
# 楷体
abstract_content_main_text_style: CharacterStyle = doc.styles.add_style("摘要内容正文", WD_STYLE_TYPE.CHARACTER, False) # type: ignore
abstract_content_main_text_style.priority = 99
abstract_content_main_text_style.quick_style = True
abstract_content_main_text_style.hidden = False
set_font_face(abstract_content_main_text_style, CNFontFaces.楷体, 'Times New Roman')

# 英文摘要标题
# 小三号Times New Roman加粗，居中，1.5倍行距，段前1行、段后1行
en_abstract_title_style: ParagraphStyle = doc.styles.add_style("英文摘要标题", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
en_abstract_title_style.priority = 99
en_abstract_title_style.quick_style = True
en_abstract_title_style.hidden = False
en_abstract_title_style.font.size = Pt(FontSizes.小三)
set_font_face(en_abstract_title_style, CNFontFaces.宋体, 'Times New Roman')
set_bold(en_abstract_title_style)
en_abstract_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
en_abstract_title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
en_abstract_title_style.paragraph_format.space_before = Pt(FontSizes.小三)
en_abstract_title_style.paragraph_format.space_after = Pt(FontSizes.小三)
en_abstract_title_style.paragraph_format.first_line_indent = 0

# 英文摘要内容
# 小四号Times New Roman加粗并加方括号，左边顶格
en_abstract_content_style: ParagraphStyle = doc.styles.add_style("英文摘要内容", WD_STYLE_TYPE.PARAGRAPH, False) # type: ignore
en_abstract_content_style.priority = 99
en_abstract_content_style.quick_style = True
en_abstract_content_style.hidden = False
en_abstract_content_style.font.size = Pt(FontSizes.小四)
en_abstract_content_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
en_abstract_content_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
en_abstract_content_style.paragraph_format.space_before = 0
en_abstract_content_style.paragraph_format.space_after = 0
en_abstract_content_style.paragraph_format.first_line_indent = 0

# 英文摘要内容标签
# Times New Roman
en_abstract_content_tag_style: CharacterStyle = doc.styles.add_style("英文摘要内容标签", WD_STYLE_TYPE.CHARACTER, False) # type: ignore
en_abstract_content_tag_style.quick_style = True
en_abstract_content_tag_style.hidden = False
en_abstract_content_tag_style.priority = 99
set_bold(en_abstract_content_tag_style)
set_font_face_en(en_abstract_content_tag_style, 'Times New Roman')

# 代码样式
# 代码格式为五号，英文为Times New Roman；中文为宋体，单倍行距。
table_content: ParagraphStyle = doc.styles.add_style("表格内容", WD_STYLE_TYPE.PARAGRAPH, True) # type: ignore
table_content.priority = 99
table_content.quick_style = True
table_content.hidden = False
table_content.font.size = Pt(FontSizes.五号)
set_font_face(table_content, CNFontFaces.宋体, 'Times New Roman')
table_content.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_content.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
table_content.paragraph_format.space_before = 0
table_content.paragraph_format.space_after = 0


with open(src_path, "r+", -1, "utf8") as src_file:
  src_content = src_file.read()
  src_lines = src_content.split("\n")
  # tokens = md.parse(src_content.replace("\n", "\n\n"), {})
  tokens = md.parse(src_content, {})
  # print(tokens)

  # 预先加载元信息
  for ti in range(len(tokens)):
    token = tokens[ti]
    if token.type == "front_matter":
      meta_data.update(yaml_load(token.content, yaml_Loader))

print(f"[操作] 合并模板 {path.abspath(template_path)}。")
template_doc: DocumentType = Document(template_path)
template_doc_el = template_doc._body._body

# 将中山大学新华学院 LOGO 调整到合适的位置
voffset_el: xmlchemy.BaseOxmlElement = template_doc_el.xpath('./w:p[@w14:textId="534624F7"]/w:r/w:drawing/wp:anchor/wp:positionV/wp:posOffset')[0]
voffset_el.text = "-0" # type: ignore
hoffset_el: xmlchemy.BaseOxmlElement = template_doc_el.xpath('./w:p[@w14:textId="534624F7"]/w:r/w:drawing/wp:anchor/wp:positionH/wp:posOffset')[0]
hoffset_el.text = "300000" # type: ignore

class MustachesType:
  cover_main = "cover_main"
  cover_title = "cover_title"
  cover = "cover"
  ais_date = "ais_date"
  abstract_title = "abstract_title"
  abstract_content = "abstract_content"
  en_abstract_title = "en_abstract_title"
  en_abstract_content = "en_abstract_content"

pi_offset = 0
for _pi in range(template_doc.paragraphs.__len__()):
  pi = _pi + pi_offset
  p = template_doc.paragraphs[_pi + pi_offset]
  mustaches = {
    "cover_title": ["论文名称", MustachesType.cover_title],
    "cover_department": ["院系", MustachesType.cover_main],
    "cover_major": ["专业", MustachesType.cover_main],
    "cover_student_name": ["学生姓名", MustachesType.cover_main],
    "cover_student_id": ["学号", MustachesType.cover_main],
    "cover_instructor": ["指导老师", MustachesType.cover_main],
    "cover_year": ["年份", MustachesType.cover],
    "cover_month": ["月份", MustachesType.cover], 
    "ais_date": ["学术诚信声明日期", MustachesType.ais_date],
    "ais_date2": ["学术诚信声明日期", MustachesType.ais_date],
    "abstract_title": ["论文名称", MustachesType.abstract_title],
    "abstract_content": ["摘要内容", MustachesType.abstract_content],
    "abstract_keywords": ["摘要关键词", MustachesType.abstract_content],
    "en_abstract_title": ["英文摘要标题", MustachesType.en_abstract_title],
    "en_abstract_content": ["英文摘要内容", MustachesType.en_abstract_content],
    "en_abstract_keywords": ["英文摘要关键词", MustachesType.en_abstract_content],
  }
  p_content = p.text.strip()
  if p_content.__len__() > 0:
    if p_content.startswith(
      "本人所呈交的毕业论文（设计），是在导师的指导下，独立进行研究工作所取得的成果，所有数据、图片资料均真实可靠。"
      ) or re.match(r"本人签名：\s*指导教师签名：", p_content):
      p.style = ais_text_style
    for mustache_name, metadata_key in mustaches.items():
      matched = re.search("{{"+mustache_name+"}}", p.text, re.MULTILINE)
      if not matched:
        continue
      before = p.text[0:matched.start()].strip()
      after = p.text[matched.end():]
      if isinstance(metadata_key, list):
        key, m_type = metadata_key
        p.clear()
        if m_type == MustachesType.cover_main or m_type == MustachesType.cover_title:
          MAX_SPACE_WIDTH = 72
          before = before.strip().replace(" ", "\u00a0")
          before_width = calculate_char_map_space_sum(count_char_map(before))

          content = get_meta_data_key(key)
          content_width = calculate_char_map_space_sum(count_char_map(content))

          p.style = cover_text_style
          p.add_run(before)
          remain = MAX_SPACE_WIDTH - before_width - content_width

          def get_space_num(remain: float):
            remain_half = math.floor(remain) // 2
            offset = 1 if remain % 1 > 0.5 else 0
            first_offset = 1 if remain % 1 > 0.75 else 0
            return remain_half + offset + first_offset, remain_half + offset

          if m_type == MustachesType.cover_title:
            MIN_SPACE_REMAIN_SPACE_WIDTH = 10
            sl = template_doc.paragraphs[pi + 1].insert_paragraph_before("", cover_text_style)
            pi_offset += 1
            fl_content_width_sum = 0
            if remain < MIN_SPACE_REMAIN_SPACE_WIDTH:
              fl_remain_without_before = 68 - before_width
              content_widths = get_count_char_space_width(content)
              fl_content = ""
              for i in range(content_widths.__len__()):
                fl_content_width_sum += content_widths[i]
                fl_content += content[i]
                if fl_remain_without_before - fl_content_width_sum < MIN_SPACE_REMAIN_SPACE_WIDTH:
                  break
              fl_remain = fl_remain_without_before - fl_content_width_sum
              fll, flr = get_space_num(fl_remain)
              p.add_run("\u00a0" * fll + fl_content  + "\u00a0" * flr, cover_text_content_style)

              sl_content_remain = content_width - fl_content_width_sum
              sl.add_run("\u00a0" * round(before_width))
              sll, slr = get_space_num(MAX_SPACE_WIDTH - sl_content_remain - before_width)
              sl.add_run("\u00a0" * sll + content[fl_content.__len__():]  + "\u00a0" * slr, cover_text_content_style)
            else:
              fll, flr = get_space_num(remain)
              p.add_run("\u00a0" * fll + content  + "\u00a0" * flr, cover_text_content_style)

              sl_content_remain = content_width - fl_content_width_sum
              sl.add_run("\u00a0" * round(before_width))
              sll, slr = get_space_num(MAX_SPACE_WIDTH - before_width)
              sl.add_run("\u00a0" * sll + "\u00a0" * slr, cover_text_content_style)
          else:
            l, r = get_space_num(remain)
            # 普通空格 Word 不识别，这里使用 nbsp 作为空格。
            run = p.add_run("\u00a0" * l + content  + "\u00a0" * r, cover_text_content_style)  
        elif m_type == MustachesType.cover:
          p.text = before + get_meta_data_key(key) + after
          p.style = cover_text_style
        elif m_type == MustachesType.abstract_title:
          p.text = get_meta_data_key(key)
          p.style = abstract_title_style
        elif m_type == MustachesType.abstract_content:
          p.clear()
          p.add_run(before, abstract_content_tag_style)
          p.add_run(get_meta_data_key(key), abstract_content_main_text_style)
          p.style = abstract_content_style
        elif m_type == MustachesType.en_abstract_title:
          p.text = get_meta_data_key(key)
          p.style = en_abstract_title_style
        elif m_type == MustachesType.en_abstract_content:
          p.clear()
          p.add_run(before, en_abstract_content_tag_style)
          p.add_run(get_meta_data_key(key), abstract_content_main_text_style)
          p.style = en_abstract_content_style
        elif m_type == MustachesType.ais_date:
          p.style = ais_text_style
          p.text = before + get_meta_data_key(key) + after
          # if key == "ais_date2":
          #   # 防止后面的页眉影响到封面和学术诚信声明
        else:
          p.text = before + get_meta_data_key(key) + after
      else:
        p.text = before + get_meta_data_key(metadata_key) + after

for it in template_doc.styles:
  try:
    doc.styles[it.name] # type: ignore
  except:
    # print("added", it.name)
    composer.add_styles(template_doc, it._element)
    
composer.append(template_doc, False)
print("[操作] 合并模板完成。")

doc = composer.doc

# [问题修复] 自动生成目录
add_toc(doc)

# [问题修复] 这段代码解决以下问题：
# 默认的页眉宽度过宽，导致右侧（标题）文本异常溢出。
header_style_element = doc.styles["Header"].element
ppr: xmlchemy.BaseOxmlElement = header_style_element.first_child_found_in("w:pPr").first_child_found_in("w:tabs") # type: ignore
right_tab: xmlchemy.BaseOxmlElement = ppr[1]
create_attribute(right_tab, "w:pos", "8650")

# (生成 style_list_file.txt 的测试代码，应该删除)
style_list = []
for it in doc.styles:
  style_list.append(str(it.name) + ":" + it.style_id)
with open(path.join(temp_dir_path, "style_list_file.txt"), "w+", -1, "utf8") as style_list_file:
  style_list_file.write("\n".join(style_list))

src_lines = src_content.split("\n")
first_title_1 = True
last_level_1_heading_text = ""
# 最后的二级标题。每当开启新的一级标题的时候会被检查和清理。
last_level_2_heading_line = None
last_level_2_heading_text = None
# 最后的列表项的信息
last_list_item = None
last_list_item_level = None
last_list_item_type = None

def get_meta_data_paper_title():
  global meta_data
  paper_title = meta_data["论文名称"]
  if not paper_title:
    print(
      f"{bcolors.WARNING}[信息缺失] 你没有提供论文名称，请手动填写页眉的论文，或在 Markdown 文档开头使用设置 name 属性。\n{bcolors.ENDC}"
    )
    paper_title = "论文名称缺失！请在此处填写以补全。"
  return paper_title

lists = []
lists_last_p = []

# # 添加目录章节
# toc_heading = doc.add_paragraph("目录", f"Heading 1")
# add_toc()
# s = doc.sections[-1]
# toc_header = s.header
# toc_header.paragraphs[0].text = (f"{get_meta_data_name()}\t\t{last_level_1_heading_text}")
# last_level_1_heading_text = "目录"
# procress_header(toc_header)

# print(tokens)

mode = "normal"
table = None
table_row = None
table_row_cell_index = 0
table_first_row = True
for ti in range(len(tokens)):
  token = tokens[ti]

  if mode == "normal":
    list_token = None
    list_style = None
    list_type = None

    lists_len = lists.__len__()
    if lists_len > 0:
      list_token = lists[-1]
      list_style = "List Bullet"
      list_type = "bullet"
      if list_token.type == "ordered_list_open":
        list_style = "List Number"
        list_type = "order"

      if lists_len > 1:
        list_style += f" {lists_len}"

    if token.type == "heading_open":
      level = int(token.tag[1:])
      text = re.sub(r"\s", "\u00a0", tokens[ti + 1].content.strip())

      # if level > 3:
      #   print(
      #     f"{bcolors.WARNING}[警告] 使用了规范未规定的 {bcolors.BOLD}{level}{bcolors.ENDC}{bcolors.WARNING} 级标题，可能会导致论文不合规范。\n{bcolors.ENDC}"
      #     + preprint_content_in_map(src_lines, token.map) + "\n"
      #   )
      # if level <= 3 and ti - 1 > 0 and tokens[ti - 1].type == "heading_close":
      #   print(
      #     f"{bcolors.WARNING}[警告] 两级标题之间要有过渡性文字。可以通过一段话引出下面的文字或者对本章内容概括。\n{bcolors.ENDC}"
      #     + preprint_content_in_map(src_lines, [tokens[ti - 3].map[0], token.map[1]]) + "\n" # type: ignore
      #   )
      if level == 1:
        # s = doc.sections[doc.sections.__len__() - 1]
        # if first_title_1:
        #   first_title_1 = False
        # else:
        s = doc.add_section(WD_SECTION.CONTINUOUS)
        header = s.header

        if last_level_1_heading_text != "":
          header.paragraphs[0].text = (f"{get_meta_data_paper_title()}\t\t{last_level_1_heading_text}")
        procress_header(header)

        last_level_1_heading_text = text
        # if last_level_2_heading_line != None:
        #   if last_level_2_heading_text != "本章小结":
        #     print(
        #       f"{bcolors.WARNING}[警告] 除第1章和最后1章外，每章最后一节为“本章小结”。\n{bcolors.ENDC}"
        #     )
        #   last_level_2_heading_text = None
        #   last_level_2_heading_line = None

      # if level == 2:
      #   last_level_2_heading_text = text
      #   last_level_2_heading_line = token.map[0] # type: ignore

      heading = doc.add_paragraph(text, f"Heading {level}")
    elif token.type == "paragraph_open":
        inline_token = tokens[ti + 1]
        if inline_token.children:
          for tk in inline_token.children:
            if tk.type == "text":
              if last_level_1_heading_text == "参考文献":
                matched = re.search(r"\[\s*(\d+)\s*\]\s*(.+)", tk.content)
                if not matched:
                  p = doc.add_paragraph(tk.content, "参考文献")
                else:
                  p = doc.add_paragraph(f"[{matched.group(1)}]\t{matched.group(2)}", "参考文献")
                  # p.add_run(matched.group(2).strip())
                  p.paragraph_format.tab_stops.add_tab_stop(
                    Pt(FontSizes.小五 * 2.4), # type: ignore
                  )
              elif lists_len > 0:
                p = doc.add_paragraph(tk.content, list_style)
                prev = lists_last_p[-1]
                # print(prev.text if prev else None, p.text)
                list_number(doc, p, prev, level=lists_len - 1, num = list_type == "order")
                lists_last_p[-1] = p
              else:
                p = doc.add_paragraph(tk.content)
            elif tk.type == "image":
              doc.add_picture(path.join(path.dirname(src_path), tk.attrs["src"]), Pt(380)) # type: ignore
              last_paragraph = doc.paragraphs[-1] 
              last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
              doc.add_paragraph(tk.content, "插图标题")
            elif tk.type == "softbreak":
              pass
            else:
              # print(tk.type, tk.content)
              doc.add_paragraph(tk.content, list_style if list_style else None)
        else:
          text = inline_token.content
          doc.add_paragraph(text)
    elif token.type == "fence":
      if token.info == "mermaid":
        id = str(time.time_ns()) + "-" + str(get_id())
        temp_mmd_path = path.abspath(path.join(temp_dir_path, id + ".mmd"))
        temp_svg_path = path.abspath(path.join(temp_dir_path, id + ".png"))
        with open(temp_mmd_path, "w+", -1, "utf8") as temp_mmd:
          temp_mmd.write(token.content)
        requests_get(f"http://127.0.0.1:{node_server['port']}/render_mermaid?src={temp_mmd_path}&target={temp_svg_path}")
        doc.add_picture(temp_svg_path, Pt(380))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ti - 1 > 0 and tokens[ti - 1].type == "html_block":
          html_block = tokens[ti - 1]
          matched = re.match(r"<!--(?P<content>(.|\s|\n)*?)-->", html_block.content)
          if matched:
            try:
              data = json.loads(matched.group("content"))
              if data["name"]:
                doc.add_paragraph(data["name"], "插图标题")
            except: pass
      else:
        text = tokens[ti].content
        doc.add_paragraph(text, "代码")
        
    elif token.type == "bullet_list_open":
      lists.append(token)
      lists_last_p.append(None)
    elif token.type == "ordered_list_open":
      lists.append(token)
      lists_last_p.append(None)
    elif token.type == "bullet_list_close":
      lists.pop()
      lists_last_p.pop()
    elif token.type == "ordered_list_close":
      lists.pop()
      lists_last_p.pop()
    elif token.type == "table_open":
      table = doc.add_table(0, 0, "Table Grid")
      table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
      set_font_face(table.style, CNFontFaces.宋体, "Times New Roman")
      # 表格居中
      tblPr = table._tblPr
      jc = create_element("w:jc")
      create_attribute(jc, "w:val", "center")
      tblPr.append(jc)

      mode = "table"
      table_first_row = True


  elif mode == "table":
    if token.type == "thead_open":
      mode = "table_thead_open"
    elif token.type == "tr_open":
      table_row = table.add_row() # type: ignore

      if table_first_row:
        # 把第一行作为行头
        trPr = table_row._element.getchildren()[0]
        tblHeader = create_element("w:tblHeader")
        trPr.append(tblHeader)
        table_first_row = False

      table_row_cell_index = 0
    elif token.type == "td_open":
      cell = table_row.cells[table_row_cell_index] # type: ignore
      cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

      # 设置宽度自适应
      tcW = cell._element.getchildren()[0].getchildren()[0]
      create_attribute(tcW, "w:type", "auto")

      p = cell.paragraphs[0]
      p.style = table_content
      p.add_run(tokens[ti + 1].content) 
      table_row_cell_index += 1
    elif token.type == "table_close":
      p = doc.add_paragraph()
      p._element
      table.autofit = True # type: ignore
      mode = "normal"
    

  elif mode == "table_thead_open":
    if token.type == "th_open":
      col = table.add_column(Pt(100)) # type: ignore
    elif token.type == "thead_close":
      mode = "table"



s = doc.add_section(WD_SECTION.CONTINUOUS)
header = s.header
header.paragraphs[0].text = (f"{get_meta_data_paper_title()}\t\t{last_level_1_heading_text}")
doc.sections[0].header.paragraphs[0].style.paragraph_format.space_after = Pt(4) # type: ignore
procress_header(header)


# 设置每页的页码
doc.sections[0].footer.paragraphs[0].style.font.size = Pt(FontSizes.小五) # type: ignore
doc.sections[0].footer.paragraphs[0].style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
set_font_face(doc.sections[0].footer.paragraphs[0].style, "宋体", 'Times New Roman') # type: ignore

doc.sections[1].footer.is_linked_to_previous = False
pg_num_type = create_page_number_type(start_num=1, format="upperRoman")
doc.sections[1]._sectPr.append(pg_num_type)
add_page_number(doc.sections[1].footer.paragraphs[0].add_run())

pg_num_type = create_page_number_type(format="upperRoman")
doc.sections[2]._sectPr.append(pg_num_type)
# doc.sections[2].footer.is_linked_to_previous = False

pg_num_type = create_page_number_type(format="upperRoman")
doc.sections[3]._sectPr.append(pg_num_type)
# doc.sections[3].footer.is_linked_to_previous = False

doc.sections[4].footer.is_linked_to_previous = False
pg_num_type = create_page_number_type(start_num=1)
doc.sections[4]._sectPr.append(pg_num_type)
add_page_number(doc.sections[4].footer.paragraphs[0].add_run())

# 设置每节的页外边距。
# 正文页边距：上：3cm，下：2.5 cm，左：2.9cm，右：2.9 cm
sections = doc.sections
for section in sections:
  section.top_margin = Cm(3)
  section.bottom_margin = Cm(2.5)
  section.left_margin = Cm(2.9)
  section.right_margin = Cm(2.9)
  section.page_width = Cm(21)
  section.page_height = Cm(29.7)

# [问题修复] 摘要和目录缺少页眉
doc.sections[1].header.is_linked_to_previous = False
doc.sections[1].header.paragraphs[0].text = f"{get_meta_data_paper_title()}\t\t摘要"
procress_header(doc.sections[1].header)
doc.sections[2].header.is_linked_to_previous = False
doc.sections[2].header.paragraphs[0].text = f"{get_meta_data_key('英文摘要标题')}\t\tAbstract"
procress_header(doc.sections[2].header)
doc.sections[3].header.is_linked_to_previous = False
doc.sections[3].header.paragraphs[0].text = f"{get_meta_data_paper_title()}\t\t目录"
procress_header(doc.sections[3].header)

  
doc.save(f"{temp_dir_path}/main_doc.docx")
composer.save(f"{temp_dir_path}/temp.docx")

# for s in composer.doc.sections:
#   print("------")
#   print(s.header.paragraphs[0].text)
#   print("--")
#   for p in s.iter_inner_content():
#     print(p.text)
#   print(s.header._element.xml)

try:
  composer.save(target_path)

  print(f"\n\n{bcolors.OKGREEN}----生成成功----{bcolors.ENDC}\n")
  print(f"\n\n{bcolors.OKGREEN}文件储存在 {bcolors.BOLD}{path.abspath(target_path)}{bcolors.ENDC}\n")
  print(f"{bcolors.OKGREEN}{bcolors.BOLD}[重要提醒] 请在生成后的文档的“目录”章节中补充目录，之前版本不支持生成目录。{bcolors.ENDC}")
except PermissionError as e:
  print(f"\n\n{bcolors.FAIL}----发生错误----{bcolors.ENDC}\n")
  print(
    f"{bcolors.FAIL}[错误] 文件无法保存！由于文件被占用，导致生成的 Word 文件无法保存。{bcolors.ENDC}")
  print(
    f"{bcolors.FAIL}{bcolors.BOLD}[提醒] 这个 Word 文件是否已经被打开？如果是，请关闭文件后再重新运行本程序。{bcolors.ENDC}")
  print("原始报错信息：", e)
except Exception as e:
  print(f"\n\n{bcolors.FAIL}----发生错误----{bcolors.ENDC}\n")
  print(f"{bcolors.FAIL}[错误] 文件无法保存！由于未知的原因，生成的 Word 文件无法保存。{bcolors.ENDC}")
  print("原始报错信息：", e)

node_server["server"].kill()
